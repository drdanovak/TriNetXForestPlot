import io
import csv
import re
from pathlib import Path

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

plt.style.use("seaborn-v0_8-whitegrid")

st.set_page_config(
    layout="wide",
    page_title="Forest Plot Generator",
    page_icon="🌲",
    initial_sidebar_state="expanded",
)

# ── Minimal CSS polish ────────────────────────────────────────────────────────
st.markdown(
    """
    <style>
    /* Tighten sidebar padding */
    section[data-testid="stSidebar"] > div { padding-top: 1rem; }
    /* Subtle section dividers */
    hr { margin: 0.5rem 0; border-color: #e0e0e0; }
    /* Make error/warning banners less alarming */
    .stAlert { border-radius: 8px; }
    </style>
    """,
    unsafe_allow_html=True,
)

REQUIRED_COLS = ["Outcome", "Effect Size", "Lower CI", "Upper CI"]
DELETE_COL = "🗑 Delete"
ORDER_COL = "↕ Order"
HEADER_COL = "🅷 Header"

# ── Defaults (centralised so they're easy to change) ─────────────────────────
DEFAULTS = dict(
    plot_title="Forest Plot",
    x_axis_label="Effect Size (RR / OR / HR)",
    ref_line=1.0,
    ref_line_label=" ",
    point_size=10,
    line_width=2,
    font_size=12,
    label_offset=0.05,
    axis_padding=10,
    cap_height=0.18,
    fig_width=10,
    fig_height_per_row=0.70,
    fig_height_min=3.0,
    top_headroom=0.0,
    bottom_padding=1.0,
    title_pad=12,
    ci_color="#1f77b4",
    marker_color="#d62728",
)


# ─────────────────────────────────────────────────────────────────────────────
# Table helpers
# ─────────────────────────────────────────────────────────────────────────────

def _normalize_table(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    for c in REQUIRED_COLS:
        if c not in df.columns:
            df[c] = None
    if DELETE_COL not in df.columns:
        df.insert(0, DELETE_COL, False)
    if HEADER_COL not in df.columns:
        df.insert(0, HEADER_COL, False)
    if ORDER_COL not in df.columns:
        df.insert(0, ORDER_COL, list(range(1, len(df) + 1)))

    df[DELETE_COL] = df[DELETE_COL].fillna(False).astype(bool)
    df[HEADER_COL] = df[HEADER_COL].fillna(False).astype(bool)
    df[ORDER_COL] = pd.to_numeric(df[ORDER_COL], errors="coerce")

    if df[ORDER_COL].isna().any():
        max_order = df[ORDER_COL].max(skipna=True)
        max_order = 0.0 if pd.isna(max_order) else float(max_order)
        for k, idx in enumerate(df.index[df[ORDER_COL].isna()].tolist(), start=1):
            df.loc[idx, ORDER_COL] = max_order + k

    df = df.sort_values(ORDER_COL, kind="mergesort").reset_index(drop=True)
    df[ORDER_COL] = range(1, len(df) + 1)

    hdr = df[HEADER_COL].fillna(False).astype(bool)
    if hdr.any():
        df.loc[hdr, ["Effect Size", "Lower CI", "Upper CI"]] = None
    return df


def _blank_row_for(df: pd.DataFrame) -> dict:
    row = {}
    for c in df.columns:
        if c == DELETE_COL:
            row[c] = False
        elif c == HEADER_COL:
            row[c] = False
        elif c == ORDER_COL:
            row[c] = None
        elif c in ["Effect Size", "Lower CI", "Upper CI"]:
            row[c] = None
        else:
            row[c] = ""
    return row


def editable_table_with_row_ops(
    df_seed: pd.DataFrame,
    state_key: str,
    tools_expanded_default: bool = False,
):
    if state_key not in st.session_state:
        st.session_state[state_key] = _normalize_table(df_seed)

    st.session_state[state_key] = _normalize_table(st.session_state[state_key])
    df_now = st.session_state[state_key]

    st.caption(
        "Tip: Check **🅷 Header** to make a row a section header (numeric columns are ignored). "
        "You can also prefix a label with `##` for the same effect."
    )

    edited = st.data_editor(
        df_now,
        num_rows="dynamic",
        use_container_width=True,
        key=f"editor_{state_key}",
        column_config={
            ORDER_COL: st.column_config.NumberColumn(ORDER_COL, help="Row order.", step=1),
            HEADER_COL: st.column_config.CheckboxColumn(
                HEADER_COL, help="Treat this row as a section header."
            ),
            DELETE_COL: st.column_config.CheckboxColumn(
                DELETE_COL, help="Mark rows for deletion."
            ),
            "Effect Size": st.column_config.NumberColumn("Effect Size", step=0.01, format="%.4f"),
            "Lower CI": st.column_config.NumberColumn("Lower CI", step=0.01, format="%.4f"),
            "Upper CI": st.column_config.NumberColumn("Upper CI", step=0.01, format="%.4f"),
        },
    )

    st.session_state[state_key] = _normalize_table(edited)
    df_now = st.session_state[state_key]

    # ── Table export ────────────────────────────────────────────────────────
    export_cols = [c for c in df_now.columns if c != DELETE_COL]
    st.download_button(
        "📥 Download current table as CSV",
        data=df_now[export_cols].to_csv(index=False).encode("utf-8"),
        file_name="forest_plot_data.csv",
        mime="text/csv",
        key=f"dl_table_{state_key}",
    )

    if len(df_now) == 0:
        return df_now

    df_view = df_now.sort_values(ORDER_COL, kind="mergesort").reset_index(drop=True)

    with st.expander("Row tools  (move / insert / header / delete)", expanded=tools_expanded_default):
        top = st.columns([2.2, 3.8, 1.2, 1.2, 1.4, 1.4, 1.6])

        with top[0]:
            row_num = st.number_input(
                "Row #",
                min_value=1,
                max_value=len(df_view),
                value=1,
                step=1,
                key=f"rownum_{state_key}",
            )
            sel_idx = int(row_num) - 1

        with top[1]:
            o = df_view.loc[sel_idx, "Outcome"]
            o = "" if pd.isna(o) else str(o)
            is_hdr = bool(df_view.loc[sel_idx, HEADER_COL])
            es = df_view.loc[sel_idx, "Effect Size"]
            lci = df_view.loc[sel_idx, "Lower CI"]
            uci = df_view.loc[sel_idx, "Upper CI"]
            if is_hdr:
                st.write(f"Selected: **HEADER** → {o}")
            else:
                st.write(f"Selected: {o}")
                st.write(f"Effect / CI: {es}  |  {lci} – {uci}")

        def commit(df_commit: pd.DataFrame):
            st.session_state[state_key] = _normalize_table(df_commit)
            st.rerun()

        curr_order = float(pd.to_numeric(df_view.loc[sel_idx, ORDER_COL], errors="coerce"))

        with top[2]:
            if st.button("⬆️ Up", key=f"up_{state_key}", use_container_width=True, disabled=(sel_idx <= 0)):
                df2 = df_view.copy()
                ab = float(pd.to_numeric(df2.loc[sel_idx - 1, ORDER_COL], errors="coerce"))
                df2.loc[sel_idx, ORDER_COL] = ab
                df2.loc[sel_idx - 1, ORDER_COL] = curr_order
                commit(df2)

        with top[3]:
            if st.button("⬇️ Down", key=f"down_{state_key}", use_container_width=True, disabled=(sel_idx >= len(df_view) - 1)):
                df2 = df_view.copy()
                bl = float(pd.to_numeric(df2.loc[sel_idx + 1, ORDER_COL], errors="coerce"))
                df2.loc[sel_idx, ORDER_COL] = bl
                df2.loc[sel_idx + 1, ORDER_COL] = curr_order
                commit(df2)

        with top[4]:
            if st.button("➕ Above", key=f"ins_above_{state_key}", use_container_width=True):
                df2 = df_view.copy()
                nr = _blank_row_for(df2)
                nr[ORDER_COL] = curr_order - 0.5
                commit(pd.concat([df2, pd.DataFrame([nr])], ignore_index=True))

        with top[5]:
            if st.button("➕ Below", key=f"ins_below_{state_key}", use_container_width=True):
                df2 = df_view.copy()
                nr = _blank_row_for(df2)
                nr[ORDER_COL] = curr_order + 0.5
                commit(pd.concat([df2, pd.DataFrame([nr])], ignore_index=True))

        with top[6]:
            if st.button("🅷 Toggle header", key=f"toggle_hdr_{state_key}", use_container_width=True):
                df2 = df_view.copy()
                new_val = not bool(df2.loc[sel_idx, HEADER_COL])
                df2.loc[sel_idx, HEADER_COL] = new_val
                if new_val:
                    df2.loc[sel_idx, ["Effect Size", "Lower CI", "Upper CI"]] = None
                commit(df2)

        bottom = st.columns([2.0, 2.0, 3.0, 3.0])

        with bottom[0]:
            if st.button("🗑 Delete selected", key=f"del_sel_{state_key}", use_container_width=True):
                commit(df_view.drop(index=sel_idx).reset_index(drop=True))

        with bottom[1]:
            if st.button("🗑 Delete checked", key=f"del_checked_{state_key}", use_container_width=True):
                df2 = df_now.copy()
                df2 = df2.loc[~df2[DELETE_COL].fillna(False).astype(bool)].reset_index(drop=True)
                commit(df2)

        with bottom[2]:
            if st.button("✅ Clear delete checks", key=f"clear_del_{state_key}", use_container_width=True):
                df2 = df_now.copy()
                df2[DELETE_COL] = False
                commit(df2)

        with bottom[3]:
            if st.button("➕ Add header below", key=f"add_hdr_below_{state_key}", use_container_width=True):
                df2 = df_view.copy()
                nr = _blank_row_for(df2)
                nr[HEADER_COL] = True
                nr[ORDER_COL] = curr_order + 0.5
                commit(pd.concat([df2, pd.DataFrame([nr])], ignore_index=True))

    return st.session_state[state_key]


# ─────────────────────────────────────────────────────────────────────────────
# TriNetX parsing helpers (unchanged logic, same as original)
# ─────────────────────────────────────────────────────────────────────────────

def _clean_line(s) -> str:
    if s is None:
        return ""
    s = str(s).strip()
    if len(s) >= 2 and s[0] == '"' and s[-1] == '"':
        s = s[1:-1]
    return s.strip()


def _to_float(x):
    try:
        if x is None:
            return None
        x = str(x).strip()
        return None if x == "" else float(x)
    except Exception:
        return None


def _extract_ci_from_string(s: str):
    nums = re.findall(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", str(s or ""))
    if len(nums) >= 3:
        return (_to_float(nums[0]), _to_float(nums[1]), _to_float(nums[2]))
    return (None, None, None)


def _parse_section_effect(lines, section_name: str):
    results = []
    n = len(lines)
    for i, line in enumerate(lines):
        if _clean_line(line).lower() != section_name.lower():
            continue
        j = i + 1
        while j < n and _clean_line(lines[j]) == "":
            j += 1
        if j >= n:
            continue
        try:
            header = next(csv.reader([lines[j]]))
        except Exception:
            header = [lines[j]]
        header = [h.strip() for h in header if h is not None]

        k = j + 1
        while k < n and _clean_line(lines[k]) == "":
            k += 1
        if k >= n:
            continue
        try:
            vals = next(csv.reader([lines[k]]))
        except Exception:
            vals = [lines[k]]
        vals = [v.strip() for v in vals]

        lower_idx = upper_idx = eff_idx = None
        for idx, h in enumerate(header):
            hl = h.lower()
            if ("ci" in hl and "lower" in hl) or ("95" in hl and "lower" in hl):
                lower_idx = idx
            if ("ci" in hl and "upper" in hl) or ("95" in hl and "upper" in hl):
                upper_idx = idx
        for idx, h in enumerate(header):
            if section_name.lower().replace(" ", "") in h.lower().replace(" ", ""):
                eff_idx = idx
                break
        if eff_idx is None:
            eff_idx = 0

        eff = _to_float(vals[eff_idx]) if eff_idx < len(vals) else None
        lci = _to_float(vals[lower_idx]) if (lower_idx is not None and lower_idx < len(vals)) else None
        uci = _to_float(vals[upper_idx]) if (upper_idx is not None and upper_idx < len(vals)) else None

        if eff is None or lci is None or uci is None:
            eff2, lci2, uci2 = _extract_ci_from_string(" ".join(vals))
            eff = eff if eff is not None else eff2
            lci = lci if lci is not None else lci2
            uci = uci if uci is not None else uci2

        if eff is not None and lci is not None and uci is not None:
            results.append({"Effect Type": section_name, "Effect Size": eff, "Lower CI": lci, "Upper CI": uci})
    return results


def parse_trinetx_export_text(text: str, filename: str) -> pd.DataFrame:
    lines = (text or "").splitlines()
    if lines:
        lines[0] = lines[0].lstrip("\ufeff")
    title = next((l for l in lines if _clean_line(l)), None)
    base_outcome = Path(filename).stem
    extracted = []
    for section in ["Risk Ratio", "Hazard Ratio", "Odds Ratio"]:
        extracted.extend(_parse_section_effect(lines, section))
    if not extracted:
        return pd.DataFrame()
    df = pd.DataFrame(extracted)
    df.insert(0, "Outcome", base_outcome)
    df["Source"] = title or filename
    df["File"] = filename
    return df


def parse_uploaded_trinetx_file(uploaded_file) -> pd.DataFrame:
    name = uploaded_file.name
    ext = name.split(".")[-1].lower()
    if ext == "csv":
        raw = uploaded_file.getvalue().decode("utf-8-sig", errors="replace")
        return parse_trinetx_export_text(raw, name)
    if ext == "xlsx":
        xls = pd.ExcelFile(uploaded_file)
        flat_lines = []
        for sheet in xls.sheet_names:
            sdf = pd.read_excel(xls, sheet_name=sheet, header=None)
            for _, r in sdf.iterrows():
                vals = [str(v).strip() for v in r.tolist() if pd.notnull(v) and str(v).strip() != ""]
                if vals:
                    flat_lines.append(",".join(vals))
        return parse_trinetx_export_text("\n".join(flat_lines), name)
    if ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(uploaded_file.getvalue()))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append(",".join(cells))
        return parse_trinetx_export_text("\n".join(lines), name)
    return pd.DataFrame()


def insert_section_headers(df: pd.DataFrame, group_col: str) -> pd.DataFrame:
    if df.empty or group_col not in df.columns:
        return df
    out_rows = []
    for g, sub in df.groupby(group_col, sort=False):
        out_rows.append({"Outcome": str(g), "Effect Size": None, "Lower CI": None, "Upper CI": None, HEADER_COL: True})
        out_rows.extend(sub[REQUIRED_COLS].to_dict("records"))
    return pd.DataFrame(out_rows)


# ─────────────────────────────────────────────────────────────────────────────
# Paste-from-clipboard helper
# ─────────────────────────────────────────────────────────────────────────────

def parse_pasted_text(text: str) -> pd.DataFrame:
    """
    Accept tab- or comma-separated text pasted from Excel / Google Sheets.
    Expects at minimum columns: Outcome, Effect Size, Lower CI, Upper CI
    (case-insensitive, order-independent).  A 'header' column is optional.
    """
    text = text.strip()
    if not text:
        return pd.DataFrame()

    sep = "\t" if "\t" in text else ","
    try:
        df = pd.read_csv(io.StringIO(text), sep=sep, dtype=str)
    except Exception:
        return pd.DataFrame()

    # Normalise column names
    rename = {}
    for col in df.columns:
        cl = col.strip().lower()
        if cl == "outcome":
            rename[col] = "Outcome"
        elif cl in ("effect size", "effectsize", "es", "estimate"):
            rename[col] = "Effect Size"
        elif "lower" in cl:
            rename[col] = "Lower CI"
        elif "upper" in cl:
            rename[col] = "Upper CI"
        elif "header" in cl:
            rename[col] = HEADER_COL
    df = df.rename(columns=rename)

    missing = [c for c in REQUIRED_COLS if c not in df.columns]
    if missing:
        return pd.DataFrame()

    for c in ["Effect Size", "Lower CI", "Upper CI"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")

    return df[REQUIRED_COLS].copy()


# ─────────────────────────────────────────────────────────────────────────────
# Plot builder  (extracted from the button handler so it can be called freely)
# ─────────────────────────────────────────────────────────────────────────────

def build_forest_plot(plot_df: pd.DataFrame, cfg: dict) -> plt.Figure:
    """
    Build and return a matplotlib Figure for the forest plot.
    cfg must contain all settings that were previously spread across the sidebar.
    """
    indent = "\u00A0" * 4
    rows, y_labels, text_styles = [], [], []
    group_mode = False

    for _, row in plot_df.iterrows():
        outcome_val = row.get("Outcome", "")
        outcome_val = "" if pd.isna(outcome_val) else str(outcome_val)

        is_header = bool(row.get(HEADER_COL, False))
        if cfg["use_groups"] and outcome_val.startswith("##"):
            is_header = True

        if is_header:
            header_txt = outcome_val.lstrip("#").strip()
            y_labels.append(header_txt)
            text_styles.append("bold")
            rows.append(None)
            group_mode = True
        else:
            display_name = f"{indent}{outcome_val}" if group_mode else outcome_val
            y_labels.append(display_name)
            text_styles.append("normal")
            rows.append(row)

    n = len(y_labels)
    fig_h = max(cfg["fig_height_min"], n * cfg["fig_height_per_row"])
    fig, ax = plt.subplots(figsize=(cfg["fig_width"], fig_h))

    # ── X-axis limits ────────────────────────────────────────────────────────
    if cfg["use_custom_xlim"]:
        if cfg["x_end"] <= cfg["x_start"]:
            raise ValueError("Custom X-axis end must be greater than start.")
        if cfg["use_log"] and (cfg["x_start"] <= 0 or cfg["x_end"] <= 0):
            raise ValueError("Log scale requires positive X-axis limits.")
        ax.set_xlim(cfg["x_start"], cfg["x_end"])
    else:
        ci_series = pd.concat(
            [plot_df["Lower CI"].dropna(), plot_df["Upper CI"].dropna()], ignore_index=True
        )
        if ci_series.empty:
            raise ValueError("No valid CI values found. Please check your data.")
        if cfg["use_log"]:
            ci_series = ci_series[ci_series > 0]
            if ci_series.empty:
                raise ValueError("Log scale requires positive CI values.")
        x_min, x_max = ci_series.min(), ci_series.max()
        if x_min == x_max:
            x_min *= 0.9
            x_max *= 1.1
        pad = (x_max - x_min) * (cfg["axis_padding"] / 100)
        ax.set_xlim(x_min - pad, x_max + pad)

    # ── Data rows ────────────────────────────────────────────────────────────
    for i, row in enumerate(rows):
        if row is None:
            continue
        effect = row["Effect Size"]
        lci = row["Lower CI"]
        uci = row["Upper CI"]
        if pd.notnull(effect) and pd.notnull(lci) and pd.notnull(uci):
            ax.hlines(i, xmin=lci, xmax=uci, color=cfg["ci_color"], linewidth=cfg["line_width"], capstyle="round")
            ax.vlines([lci, uci], i - cfg["cap_height"], i + cfg["cap_height"], color=cfg["ci_color"], linewidth=cfg["line_width"])
            ax.plot(effect, i, "o", color=cfg["marker_color"], markersize=cfg["point_size"])
            if cfg["show_values"]:
                label = f"{effect:.2f} [{lci:.2f}, {uci:.2f}]"
                ax.text(
                    uci + cfg["label_offset"],
                    i,
                    label,
                    va="center",
                    fontsize=max(8, cfg["font_size"] - 2),
                )

    # ── Reference line ───────────────────────────────────────────────────────
    ref = cfg["ref_line"]
    if ref is not None:
        ax.axvline(x=ref, color="gray", linestyle="--", linewidth=1.2, zorder=0)
        if cfg["ref_line_label"]:
            ax.text(
                ref,
                -1 - cfg["top_headroom"] + 0.3,
                cfg["ref_line_label"],
                ha="center",
                va="bottom",
                fontsize=max(8, cfg["font_size"] - 3),
                color="gray",
                style="italic",
            )

    # ── Axes cosmetics ───────────────────────────────────────────────────────
    ax.set_yticks(range(n))
    tick_labels = ax.set_yticklabels(y_labels)
    for tl, style in zip(tick_labels, text_styles):
        tl.set_fontweight("bold" if style == "bold" else "normal")
        tl.set_fontsize(cfg["font_size"])

    if cfg["use_log"]:
        ax.set_xscale("log")

    ax.grid(cfg["show_grid"], axis="x", linestyle=":", linewidth=0.6)

    # top_headroom pushes the axis ceiling up (positive = more space at top of figure)
    ax.set_ylim(n - 1 + cfg["bottom_padding"], -1 - cfg["top_headroom"])
    ax.set_xlabel(cfg["x_axis_label"], fontsize=cfg["font_size"])
    ax.set_title(cfg["plot_title"], fontsize=cfg["font_size"] + 2, weight="bold", pad=cfg["title_pad"])
    ax.tick_params(axis="x", labelsize=cfg["font_size"] - 1)

    fig.tight_layout()
    return fig


# ─────────────────────────────────────────────────────────────────────────────
# Header
# ─────────────────────────────────────────────────────────────────────────────

col_title, col_live = st.columns([5, 2])
with col_title:
    st.title("🌲 Forest Plot Generator")
with col_live:
    st.write("")
    live_preview = st.toggle(
        "⚡ Live preview",
        value=False,
        help="When ON, the plot updates automatically whenever you change the data or settings. "
             "Turn OFF for large tables to avoid lag.",
    )

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# Input mode
# ─────────────────────────────────────────────────────────────────────────────

input_mode = st.radio(
    "Data input method:",
    ["✍️ Manual entry", "📋 Paste from clipboard", "📄 Import TriNetX tables", "📤 Upload structured file"],
    index=0,
    horizontal=True,
)

df = None

# ── Manual entry ─────────────────────────────────────────────────────────────
if input_mode == "✍️ Manual entry":
    default_data = pd.DataFrame(
        {
            "Outcome": ["Cardiovascular", "Hypertension", "Stroke", "Metabolic", "Diabetes", "Obesity"],
            HEADER_COL: [True, False, False, True, False, False],
            "Effect Size": [None, 1.5, 1.2, None, 0.85, 1.2],
            "Lower CI": [None, 1.2, 1.0, None, 0.7, 1.0],
            "Upper CI": [None, 1.8, 1.5, None, 1.0, 1.4],
        }
    )
    df = editable_table_with_row_ops(default_data, "manual_table_df")

# ── Paste from clipboard ─────────────────────────────────────────────────────
elif input_mode == "📋 Paste from clipboard":
    st.markdown(
        "Paste tab-separated or comma-separated data (e.g. copied from Excel or Google Sheets). "
        "The first row must be a header row with columns **Outcome**, **Effect Size**, **Lower CI**, **Upper CI** "
        "(names are case-insensitive)."
    )
    pasted = st.text_area("Paste data here:", height=180, placeholder="Outcome\tEffect Size\tLower CI\tUpper CI\nHypertension\t1.50\t1.20\t1.80\n...")
    if pasted.strip():
        parsed_paste = parse_pasted_text(pasted)
        if parsed_paste.empty:
            st.error(
                "Could not parse the pasted data. Make sure the header row contains at least: "
                "Outcome, Effect Size, Lower CI, Upper CI."
            )
        else:
            st.success(f"Parsed {len(parsed_paste)} data row(s).")
            df = editable_table_with_row_ops(parsed_paste, "paste_table_df")
    else:
        st.info("Paste your data above to get started.")

# ── TriNetX import ────────────────────────────────────────────────────────────
elif input_mode == "📄 Import TriNetX tables":
    uploaded_files = st.file_uploader(
        "Upload one or more TriNetX export tables (CSV / XLSX / DOCX)",
        type=["csv", "xlsx", "docx"],
        accept_multiple_files=True,
    )
    if uploaded_files:
        parsed_frames, failures = [], []
        for f in uploaded_files:
            try:
                p = parse_uploaded_trinetx_file(f)
                if not p.empty:
                    parsed_frames.append(p)
                else:
                    failures.append((f.name, "No Risk Ratio / Hazard Ratio / Odds Ratio section found."))
            except Exception as e:
                failures.append((f.name, str(e)))

        if failures:
            with st.expander("⚠️ Parsing warnings", expanded=False):
                for fn, msg in failures:
                    st.write(f"- **{fn}**: {msg}")

        if parsed_frames:
            parsed = pd.concat(parsed_frames, ignore_index=True)
            effect_types = sorted(parsed["Effect Type"].unique().tolist())
            default_keep = [t for t in ["Risk Ratio", "Hazard Ratio"] if t in effect_types] or effect_types
            keep_types = st.multiselect("Effect types to include:", options=effect_types, default=default_keep)
            plot_base = parsed[parsed["Effect Type"].isin(keep_types)].copy()

            col_a, col_b = st.columns(2)
            with col_a:
                append_type = st.checkbox(
                    "Append effect type to label when outcome has multiple types", value=True
                )
            with col_b:
                add_headers = st.checkbox("Insert section headers per table", value=False)

            if append_type and not plot_base.empty:
                multi = plot_base.groupby("Outcome")["Effect Type"].nunique()
                multi_outcomes = set(multi[multi > 1].index)
                plot_base["Outcome"] = plot_base.apply(
                    lambda r: f"{r['Outcome']} ({r['Effect Type']})" if r["Outcome"] in multi_outcomes else r["Outcome"],
                    axis=1,
                )

            for c in REQUIRED_COLS:
                if c not in plot_base.columns:
                    plot_base[c] = None
            for c in ["Effect Size", "Lower CI", "Upper CI"]:
                plot_base[c] = pd.to_numeric(plot_base[c], errors="coerce")

            if add_headers:
                header_grouping = st.selectbox("Header grouping field", ["Source", "File"], index=0)
                df_for_editor = insert_section_headers(plot_base, group_col=header_grouping)
            else:
                df_for_editor = plot_base[REQUIRED_COLS].copy()

            df = editable_table_with_row_ops(df_for_editor, "trinetx_import_table_df")

            with st.expander("📋 Full extracted data", expanded=False):
                st.dataframe(parsed, use_container_width=True)

            st.download_button(
                "📥 Download parsed rows as CSV",
                data=plot_base.to_csv(index=False).encode("utf-8"),
                file_name="parsed_trinetx_effects.csv",
                mime="text/csv",
            )
        else:
            st.info("No parsable TriNetX effect sections found in the uploaded files.")
    else:
        st.info("Upload one or more TriNetX export tables to get started.")

# ── Structured file upload ────────────────────────────────────────────────────
else:
    st.subheader("Upload a structured file")
    template_df = pd.DataFrame(
        {
            "Outcome": ["Cardiovascular", "Hypertension", "Stroke"],
            HEADER_COL: [True, False, False],
            "Effect Size": [None, 1.50, 1.20],
            "Lower CI": [None, 1.20, 1.00],
            "Upper CI": [None, 1.80, 1.50],
        }
    )
    st.download_button(
        "📥 Download CSV template",
        data=template_df.to_csv(index=False).encode("utf-8"),
        file_name="forest_plot_template.csv",
        mime="text/csv",
    )
    uploaded_file = st.file_uploader("Upload CSV or Excel file", type=["csv", "xlsx"])
    if uploaded_file:
        try:
            sig = (uploaded_file.name, uploaded_file.size)
            if st.session_state.get("structured_file_sig") != sig:
                st.session_state["structured_file_sig"] = sig
                st.session_state.pop("structured_table_df", None)
            if uploaded_file.name.endswith(".csv"):
                df_loaded = pd.read_csv(uploaded_file)
            else:
                df_loaded = pd.read_excel(uploaded_file)
            if not all(c in df_loaded.columns for c in REQUIRED_COLS):
                st.error(f"File must include columns: {REQUIRED_COLS}")
            else:
                df = editable_table_with_row_ops(df_loaded, "structured_table_df")
        except Exception as e:
            st.error(f"Error reading file: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Sidebar: Plot settings  (organised into clear sections)
# ─────────────────────────────────────────────────────────────────────────────

st.sidebar.header("⚙️ Plot Settings")

# ─── Labels ──────────────────────────────────────────────────────────────────
with st.sidebar.expander("🏷️ Labels", expanded=True):
    plot_title = st.text_input("Plot title", value=DEFAULTS["plot_title"])
    x_axis_label = st.text_input("X-axis label", value=DEFAULTS["x_axis_label"])
    show_grid = st.checkbox("Show grid", value=True)
    show_values = st.checkbox("Show numerical annotations on plot", value=False)
    use_groups = st.checkbox("Treat '##' prefix as section header", value=True)

# ─── Reference line ───────────────────────────────────────────────────────────
with st.sidebar.expander("📍 Reference line", expanded=True):
    ref_col1, ref_col2 = st.columns(2)
    with ref_col1:
        ref_line_val = st.number_input(
            "Reference line X", value=DEFAULTS["ref_line"], step=0.1,
            help="Set to the null value for your effect measure (e.g. 1.0 for RR/OR/HR, 0.0 for mean diff)."
        )
    with ref_col2:
        ref_line_label = st.text_input("Label", value=DEFAULTS["ref_line_label"])
    show_ref_line = st.checkbox("Show reference line", value=True)

# ─── X-axis ───────────────────────────────────────────────────────────────────
with st.sidebar.expander("📏 X-axis", expanded=False):
    use_log = st.checkbox("Log scale", value=False)
    axis_padding = st.slider("Auto-range padding (%)", 2, 40, DEFAULTS["axis_padding"])
    use_custom_xlim = st.checkbox("Custom X-axis range", value=False)
    x_start = st.number_input("Start", value=0.0, step=0.1, disabled=not use_custom_xlim)
    x_end = st.number_input("End", value=3.0, step=0.1, disabled=not use_custom_xlim)

# ─── Layout ───────────────────────────────────────────────────────────────────
with st.sidebar.expander("🧱 Layout & spacing", expanded=False):
    fig_width = st.slider("Figure width (inches)", 6, 20, DEFAULTS["fig_width"])
    fig_height_per_row = st.slider(
        "Height per row (inches)", 0.3, 1.5, DEFAULTS["fig_height_per_row"], step=0.05,
        help="Scales figure height with the number of rows."
    )
    top_headroom = st.slider("Top headroom (rows)", 0.0, 6.0, DEFAULTS["top_headroom"], step=0.5)
    bottom_padding = st.slider("Bottom padding (rows)", 0.0, 6.0, DEFAULTS["bottom_padding"], step=0.5)
    title_pad = st.slider("Title padding (points)", 0, 40, DEFAULTS["title_pad"], step=2)

# ─── Visual style ─────────────────────────────────────────────────────────────
with st.sidebar.expander("🎨 Visual style", expanded=False):
    color_scheme = st.selectbox("Color scheme", ["Color", "Black & White"])
    point_size = st.slider("Marker size", 4, 20, DEFAULTS["point_size"])
    line_width = st.slider("CI line width", 1, 5, DEFAULTS["line_width"])
    font_size = st.slider("Font size", 8, 24, DEFAULTS["font_size"])
    label_offset = st.slider("Annotation offset", 0.01, 0.5, DEFAULTS["label_offset"])
    cap_height = st.slider("CI cap height", 0.05, 0.5, DEFAULTS["cap_height"], step=0.01)

    if color_scheme == "Color":
        ci_color = st.color_picker("CI line color", DEFAULTS["ci_color"])
        marker_color = st.color_picker("Marker color", DEFAULTS["marker_color"])
    else:
        ci_color = marker_color = "black"

# ─── Export ───────────────────────────────────────────────────────────────────
with st.sidebar.expander("💾 Export", expanded=False):
    export_dpi = st.slider("PNG resolution (DPI)", 72, 600, 300, step=50)
    export_format = st.selectbox("Export format", ["PNG", "SVG", "PDF"])


# ─────────────────────────────────────────────────────────────────────────────
# Assemble config dict (so build_forest_plot() receives one clean object)
# ─────────────────────────────────────────────────────────────────────────────

cfg = dict(
    plot_title=plot_title,
    x_axis_label=x_axis_label,
    ref_line=ref_line_val if show_ref_line else None,
    ref_line_label=ref_line_label,
    show_grid=show_grid,
    show_values=show_values,
    use_groups=use_groups,
    use_log=use_log,
    axis_padding=axis_padding,
    use_custom_xlim=use_custom_xlim,
    x_start=x_start,
    x_end=x_end,
    fig_width=fig_width,
    fig_height_per_row=fig_height_per_row,
    fig_height_min=DEFAULTS["fig_height_min"],
    top_headroom=top_headroom,
    bottom_padding=bottom_padding,
    title_pad=title_pad,
    point_size=point_size,
    line_width=line_width,
    font_size=font_size,
    label_offset=label_offset,
    cap_height=cap_height,
    ci_color=ci_color,
    marker_color=marker_color,
)


# ─────────────────────────────────────────────────────────────────────────────
# Plot generation section
# ─────────────────────────────────────────────────────────────────────────────

if df is not None:
    # Prepare plot_df (sorted, cleaned)
    plot_df = df.drop(columns=[DELETE_COL], errors="ignore").copy()
    if ORDER_COL in plot_df.columns:
        plot_df[ORDER_COL] = pd.to_numeric(plot_df[ORDER_COL], errors="coerce")
        plot_df = plot_df.sort_values(ORDER_COL, kind="mergesort").reset_index(drop=True)
    for c in REQUIRED_COLS:
        if c not in plot_df.columns:
            plot_df[c] = None
    if HEADER_COL not in plot_df.columns:
        plot_df[HEADER_COL] = False
    for c in ["Effect Size", "Lower CI", "Upper CI"]:
        plot_df[c] = pd.to_numeric(plot_df[c], errors="coerce")

    # Warn about rows with incomplete data (not headers)
    data_rows = plot_df[~plot_df[HEADER_COL].fillna(False).astype(bool)]
    incomplete = data_rows[
        data_rows[["Effect Size", "Lower CI", "Upper CI"]].isnull().any(axis=1)
        & data_rows["Outcome"].notna()
        & (data_rows["Outcome"].astype(str).str.strip() != "")
    ]
    if not incomplete.empty:
        with st.expander(f"⚠️ {len(incomplete)} row(s) with missing values (will be skipped in plot)", expanded=False):
            st.dataframe(incomplete[["Outcome", "Effect Size", "Lower CI", "Upper CI"]], use_container_width=True)

    st.divider()

    # ── Generate / auto-preview ─────────────────────────────────────────────
    if live_preview:
        try:
            fig = build_forest_plot(plot_df, cfg)
            st.pyplot(fig, use_container_width=False)
            plt.close(fig)
        except ValueError as e:
            st.error(str(e))
    else:
        if st.button("📊 Generate Forest Plot", type="primary", use_container_width=False):
            try:
                with st.spinner("Rendering plot…"):
                    fig = build_forest_plot(plot_df, cfg)
                st.pyplot(fig, use_container_width=False)
                st.session_state["last_fig"] = fig
            except ValueError as e:
                st.error(str(e))
        elif "last_fig" in st.session_state:
            # Show the last-generated plot between button presses so it doesn't vanish
            st.info("Showing last generated plot. Click **Generate Forest Plot** to refresh.")
            st.pyplot(st.session_state["last_fig"], use_container_width=False)

    # ── Download ─────────────────────────────────────────────────────────────
    st.markdown("#### 💾 Download plot")
    dl_col1, dl_col2, dl_col3 = st.columns(3)

    def _render_to_bytes(fmt: str, dpi: int) -> bytes:
        fig2 = build_forest_plot(plot_df, cfg)
        buf = io.BytesIO()
        fig2.savefig(buf, format=fmt.lower(), dpi=dpi if fmt == "PNG" else None, bbox_inches="tight")
        plt.close(fig2)
        return buf.getvalue()

    mime_map = {"PNG": "image/png", "SVG": "image/svg+xml", "PDF": "application/pdf"}
    ext_map = {"PNG": "png", "SVG": "svg", "PDF": "pdf"}

    with dl_col1:
        st.download_button(
            f"📥 Download as PNG",
            data=_render_to_bytes("PNG", export_dpi),
            file_name="forest_plot.png",
            mime="image/png",
        )
    with dl_col2:
        st.download_button(
            "📥 Download as SVG",
            data=_render_to_bytes("SVG", export_dpi),
            file_name="forest_plot.svg",
            mime="image/svg+xml",
        )
    with dl_col3:
        st.download_button(
            "📥 Download as PDF",
            data=_render_to_bytes("PDF", export_dpi),
            file_name="forest_plot.pdf",
            mime="application/pdf",
        )

else:
    st.info("👆 Choose a data input method above and enter your data to generate a forest plot.")
